"""python-docx extraction adapter."""
from __future__ import annotations

import docx  # type: ignore


def extract(file_path: str) -> str | None:
    try:
        document = docx.Document(file_path)
        lines = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(lines)
    except Exception:
        return None

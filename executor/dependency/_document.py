"""Shared rich-document text extraction (PDF / Word / Excel).

A single source of truth for turning a binary/zipped document into plain text,
used by BOTH the Grep tool (to search inside documents) and the Read tool (to
read them with line offsets). Keeping extraction here guarantees the two tools
agree on line numbering: a position Grep reports as ``report.pdf:42`` is the
same line Read returns for ``offset=42``.

Not a registered tool (underscore-prefixed, like ``_file_base.py``), so the
registry package scan ignores it.

Line model: every extractor returns one flat string with lines joined by
``"\\n"``. Callers split on ``"\\n"`` (see ``document_lines``) so that the i-th
line (1-indexed) is exactly the text between the (i-1)-th and i-th newline —
matching how Grep computes a match's line number.
"""
from __future__ import annotations

import os
from typing import Optional

from mote.common.const.tools import DOCUMENT_EXTENSIONS


def is_document(file_path: str) -> bool:
    """True if the path is a rich document handled via text extraction."""
    return file_path.lower().endswith(DOCUMENT_EXTENSIONS)


def document_lines(text: str) -> list[str]:
    """Split extracted text into lines the way both tools agree on.

    Uses ``str.split("\\n")`` (NOT ``splitlines()``) so line indexing matches
    Grep's ``text.count("\\n", 0, match_start) + 1`` position arithmetic exactly.
    """
    return text.split("\n")


def extract_pdf_text(file_path: str) -> Optional[str]:
    """Extract a PDF's text, pages joined by newlines.

    Tries PyMuPDF (fitz) -> pdfminer -> pypdf/PyPDF2, using whichever is
    installed. Returns the full text or None if no backend is available / the
    extraction fails.
    """
    # PyMuPDF (fitz): fastest, best layout.
    try:
        import fitz  # type: ignore

        parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                parts.append(page.get_text("text"))  # type: ignore[attr-defined]  # fitz.Page stub gap
        return "\n".join(parts)
    except ImportError:
        pass
    except Exception:
        return None

    # pdfminer.six: pure-Python, good text fidelity.
    try:
        from pdfminer.high_level import extract_text  # type: ignore

        return extract_text(file_path)
    except ImportError:
        pass
    except Exception:
        return None

    # pypdf / PyPDF2: last resort.
    try:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError:
            from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(file_path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except ImportError:
        return None
    except Exception:
        return None


def extract_docx_text(file_path: str) -> Optional[str]:
    """Extract text from a Word .docx (paragraphs + table cells) via python-docx."""
    try:
        import docx  # type: ignore
    except ImportError:
        return None
    try:
        document = docx.Document(file_path)
        lines = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells]
                lines.append("\t".join(cells))
        return "\n".join(lines)
    except Exception:
        return None


def extract_xlsx_text(file_path: str) -> Optional[str]:
    """Extract text from an Excel .xlsx via openpyxl.

    Emits one line per row as tab-joined cells, prefixed with the sheet name so
    matches can be located. read_only + data_only keeps memory bounded on large
    workbooks.
    """
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError:
        return None
    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(cells):
                    lines.append(f"[{ws.title}]\t" + "\t".join(cells))
        wb.close()
        return "\n".join(lines)
    except Exception:
        return None


def extract_document_text(file_path: str) -> Optional[str]:
    """Dispatch to the right extractor by extension. None = cannot extract."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf_text(file_path)
    if ext == ".docx":
        return extract_docx_text(file_path)
    if ext == ".xlsx":
        return extract_xlsx_text(file_path)
    return None

"""python-docx extraction adapter."""

from __future__ import annotations

import docx  # type: ignore

from mote.runtime.fileops.document_budgets import BoundedTextSink


def extract(file_path: str, *, sink: BoundedTextSink) -> None:
    document = docx.Document(file_path)
    separator = ""
    for paragraph in document.paragraphs:
        sink.write(separator)
        sink.write(paragraph.text)
        separator = "\n"
    for table in document.tables:
        for row in table.rows:
            sink.write(separator)
            sink.write("\t".join(cell.text for cell in row.cells))
            separator = "\n"
